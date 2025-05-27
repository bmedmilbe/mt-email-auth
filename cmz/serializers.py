import html
import mammoth
from pprint import pprint
from rest_framework import serializers
from django.db.transaction import atomic
from django.core.mail import BadHeaderError
from templated_mail.mail import BaseEmailMessage
from django.apps import apps
from django.conf import settings
from .models import Information, SecreatarySection, Secretary, Service, Tour, Post, ImagesTour, Messages, Post, PostDocument, PostFile, PostImages, Role, Section, Team
from core.serializers import UserCreateSerializer
import urllib3
from bs4 import BeautifulSoup
import docx2txt
import requests
from io import BytesIO
import timeago
from datetime import datetime, timezone
import re
from docx import Document
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.core.mail import get_connection, send_mail

def text_to_html_paragraphs(text):
    # First, replace multiple newlines with a single newline,
    # so you don't get empty paragraphs
    text = re.sub(r'\n\s*\n', '\n', text)

    # Split the text into lines
    lines = text.split('\n')

    # for line in lines:
    #     pprint(line)

    # Wrap each line in a <p> tag and join them
    return ''.join(f'<p>{line.strip()}</p>\n' for line in lines)

def split_html_by_paragraphs(html_string):
    """
    Splits an HTML string into a list of strings, where each string
    contains a paragraph element and its contents.

    Args:
        html_string: The HTML string to split.

    Returns:
        A list of strings, where each string is a paragraph element.
    """
    soup = BeautifulSoup(html_string, 'html.parser')
    paragraphs = soup.find_all('p')
    return [str(p) for p in paragraphs]

def addPicures(post: Post, html_paragraphs):
    full = split_html_by_paragraphs(html_paragraphs)
    new_full = list()
    for p in full:
        temp = p
        if "https://bm-edmilbe-bucket.s3" in p:
            p=p.replace("<p>", '').replace("</p>", '')
            temp = f"<img class='video-view' src='{p}' alt='{post.title}' />"

        if "https://www.youtube.com" in p:
            p=p.replace("<p>", '').replace("</p>", '')
            # temp = f"<iframe width='560' height='315' src='{p}' title='YouTube video player' frameborder='0' allow='accelerometer; autoplay; clipboard-write; encrypted-media; gyroscope; picture-in-picture; web-share' referrerpolicy='strict-origin-when-cross-origin' allowfullscreen></iframe>"

            # temp = f"<div class='embed-responsive embed-responsive-16by9'><iframe class='embed-responsive-item' src='{p}?rel=0' allowfullscreen></iframe></div>"
            temp = f"<div class='ratio ratio-16x9 video-view'><iframe src='{p}?rel=0' title='YouTube video' allowfullscreen></iframe></div>"
        new_full.append(temp)

        
            
            # pprint(''.join(new_full))
    
    return ''.join(new_full)
        
        

    return ""
    text_with_images = ""

    count = 0
    images = post.post_images.all()

    for x in html_paragraphs.split("@image"):
        if len(html_paragraphs.split("@image")) > 1 and count < len(html_paragraphs.split("@image"))-1:

            # pprint(f"<img src={post.post_images.all()[count].picture.url} alt={post.title} />")

            if len(images) > count:
                text_with_images = f"{text_with_images}{x}<img src={images[count].picture.url} alt={post.title} />"
                count = count + 1
            else:
                text_with_images = f"{text_with_images}{x}"
        elif len(html_paragraphs.split("@image")) == 1:
            text_with_images = x
        else:
            text_with_images = f"{text_with_images}{x}"

    return text_with_images
def extract_text_from_docx_bytesio(bytes_io_obj):
        """
        Extracts text from a .docx file stored in a BytesIO object.

        Args:
            bytes_io_obj: A BytesIO object containing the .docx file data.

        Returns:
            The extracted text as a string, or None if an error occurs.
        """
        try:
            document = Document(bytes_io_obj)
            full_text = []
            for paragraph in document.paragraphs:
                full_text.append(paragraph.text)
            return "\n".join(full_text)  # Join paragraphs with newlines

        except Exception as e:
            print(f"Error extracting text: {e}")
            return None
        finally:
            bytes_io_obj.seek(0) #reset the stream to the beginning.

def addVideo(post: Post, html_paragraphs):
    # pprint(len(html_paragraphs.split("@image")))
    text_with_video = ""
    count = 0
    videos = post.post_videos.all()
    # pprint(len(videos))
    for x in html_paragraphs.split("@video"):
        if len(html_paragraphs.split("@video")) > 1 and count < len(html_paragraphs.split("@video"))-1:
            if len(videos) > count:
                text_with_video = f"{text_with_video}{x}<div class='video-text embed-responsive embed-responsive-16by9'><iframe class='embed-responsive-item' src={videos[count].video} allowfullscreen></iframe></div>"
                count = count + 1
            else:
                text_with_video = f"{text_with_video}{x}"
        elif len(html_paragraphs.split("@video")) == 1:
            text_with_video = x
        else:
            text_with_video = f"{text_with_video}{x}"

    return text_with_video


# import html


def embed(link):
    return ""
    # return (

    #     '<section className="content-inner-2 wow fadeIn" data-wow-duration="2s" data-wow-delay="0.2s"><div className="container">' +

    #     '<div className="row">' +
    #     '<div className="col-lg-12">' +
    #     '<div className="video-bx style-1 overlay-black-light">' +
    #     '<img rc="https://res.cloudinary.com/dybteyiju/image/upload/v1678362568/alexander-shatov-niUkImZcSP8-unsplash_tkkhw6.jpg" alt="" />' +
    #     '                <div className="video-btn"><a href="#" className="popup-youtube" onClick={() => setOpen(true)}><i className="flaticon-play"></i></a><h2 className="title">Assista o video</h2></div>' +
    #     '             <ModalVideo channel="youtube" autoplay isOpen={isOpen} videoId="{link}" onClose={() => setOpen(false)}/></div>' +
    #     '</div>' +
    #     '</div>' +
    #     '</div>' +
    #     '</section>')

    # return f'<div class="embed-responsive embed-responsive-4by3"><iframe class="embed-responsive-item" src={link} allowfullscreen></iframe></div>'


class SectionSerializer(serializers.ModelSerializer):
    class Meta:
        model = Section
        fields = ['id', 'title']


class PostImagesSerializer(serializers.ModelSerializer):
    class Meta:
        model = PostImages
        fields = ["id", "picture"]

    def create(self, validated_data):
        post_image = validated_data
        post_image['post_id'] = self.context["post_id"]
        return super().create(post_image)


class PostSerializer(serializers.ModelSerializer):
    # owner = serializers.SerializerMethodField(
    #     method_name="get_owner")
    beginnig = serializers.SerializerMethodField(
        method_name="get_beginnig")

    posted_at = serializers.SerializerMethodField(
        method_name="get_posted_at")

    text = serializers.SerializerMethodField(
        method_name="get_text")

    post_images = PostImagesSerializer(many=True)

    class Meta:
        model = Post
        fields = ['id',
                  'title',
                  'slug',
                  'picture',
                  #   'doctor',
                  'text_file',
                  'post_images',
                  'beginnig',
                  'text',
                  'posted_at',
                  'is_cmz_service',
                  'is_social_service',
                  'date',
                  ]
 
    def get_posted_at(self, post: Post):
        now = datetime.now(timezone.utc)
        return timeago.format(post.date, now)

    def get_beginnig(self, post: Post):
        # return ""
        url = post.text_file.url
        # url = "http://127.0.0.1:8000/media/camaramz/posts/documents/text_cecab_rubish_1.docx"
        docx = BytesIO(requests.get(url).content)

        # extract text
        text = docx2txt.process(docx).replace('\n','')

        return f"{text[:100]}..."
        with open(path.data, "rb") as docx_file:

            result = mammoth.extract_raw_text(docx_file)
            return result.value.replace('\n', '').strip()  # The raw text
    

    def get_text(self, post: Post):
        # pprint( post.text_file)
        # try:
        #     response = requests.get( post.text_file.url, stream=True)
        #     response.raise_for_status() # Raise HTTPError for bad responses (4xx or 5xx)

        #     bytes_io_obj = BytesIO(response.content)
        #     document = Document(bytes_io_obj)
        #     full_text = []
        #     for paragraph in document.paragraphs:
        #         full_text.append(paragraph.text)
        #     return "\n".join(full_text)

        # except requests.exceptions.RequestException as e:
        #     print(f"Error fetching URL: {e}")
        #     return None
        # except Exception as e:
        #     print(f"Error reading docx file: {e}")
        #     return None
        # finally:
        #     if 'bytes_io_obj' in locals():
        #         bytes_io_obj.seek(0) #reset the stream to the beginning.

        # return ''
        url = post.text_file.url
        # pprint(post)
        # pprint(url)
        # url = "http://127.0.0.1:8000/media/camaramz/posts/documents/text_cecab_rubish_1.docx"

        docx = BytesIO(requests.get(url).content)
        # pprint(docx)
        # extract text
        text = docx2txt.process(docx)
        # text = extract_text_from_docx_bytesio(docx)
        # pprint(text)

        html_paragraphs = text_to_html_paragraphs(text).replace("\n", "")

        text_with_media = f"{addPicures(post, html_paragraphs)}"
        # text_with_media = f"{addVideo(post, text_with_media)}"

        return text_with_media


class RoleSerializer(serializers.ModelSerializer):
    class Meta:
        model = Role
        fields = ['id', 'title']


class TeamSerializer(serializers.ModelSerializer):

    role = RoleSerializer()

    class Meta:
        model = Team
        fields = ['id', 'role',  'name', "image"]


class MessagesSerializer(serializers.ModelSerializer):
    class Meta:
        model = Messages
        fields = ['name', 'email', 'subject', 'text']

    def create(self, validated_data):
        
        try:
            email = validated_data['email']
            parthner = parthner = "CMZ"
            convert_to_html_content =  render_to_string(
                                        template_name='emails/message.html',
                                        context={'message': validated_data['text'],
                         'section': validated_data['subject'],
                         'name': validated_data['name'], 'email': validated_data['email']
                                                , 'logo':f"{settings.EMAILS[parthner]['LOGO']}"
                                                }
                                        )
            plain_message = strip_tags(convert_to_html_content) 



            connection = get_connection(
                host=settings.EMAIL_HOST,
                port=settings.EMAIL_PORT,
                username=settings.EMAILS[parthner]['EMAIL'],
                password=settings.EMAILS[parthner]['PASSWORD'],
                use_tls=settings.EMAIL_USE_TLS)

            


            # Send an email using the custom connection
            send_mail(
                # 'Subject',
                #        msg_plain,settings.EMAILS[parthner]['EMAIL'], 
                #        [email],
                #        html_message=msg_html,  
                #     #    context={'username': 'John'},
                    
                #       connection=connection
                    
                    
                                
                subject=validated_data['subject'],
                message=plain_message,
                from_email=settings.EMAILS[parthner]['EMAIL'],
                recipient_list=[email],  
                html_message=convert_to_html_content,
                fail_silently=True,   # Optional

                    connection=connection

                    
                    )













            
            # validated_data['sent'] = True
            # message = BaseEmailMessage(
            #     template_name='emails/message.html',
            #     context={'message': validated_data['text'],
            #              'section': validated_data['subject'],
            #              'name': validated_data['name'], 'email': validated_data['email']})
            # message.send([user.email for user in users])
            return super().create(validated_data)

        except BadHeaderError:
            pass



class ImagesTourSerializer(serializers.ModelSerializer):
    class Meta:
        model = ImagesTour
        fields = ["id", "image"]


class InformationSerializer(serializers.ModelSerializer):
    class Meta:
        model = Information
        fields = ["id", "service", "question", "information"]


class ServiceSerializer(serializers.ModelSerializer):
    informations = InformationSerializer(many=True)

    class Meta:
        model = Service
        fields = ["id", "name", "slug",
                  "informations",  "picture", "description"]


class TourSerializer(serializers.ModelSerializer):

    images = ImagesTourSerializer(many=True)
    posted_at = serializers.SerializerMethodField(
        method_name="get_posted_at")

    class Meta:
        model = Tour
        fields = ["id", "title", "description",
                  "date", "images", "slug", "posted_at", "location"]

    def get_posted_at(self, post: Post):
        now = datetime.now(timezone.utc)
        return timeago.format(post.date, now)
